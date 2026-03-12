import json
import sys
import os
from datetime import datetime
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_huggingface import HuggingFaceEmbeddings
from langchain_community.vectorstores import FAISS
from langchain_community.llms import Ollama
from langchain_core.documents import Document
from docx import Document as WordDocument

# Prevent output buffering
sys.stdout.reconfigure(line_buffering=True)

DATA_PATH = "project_knowledge.json"
VECTOR_DB_PATH = "project_vector_db"

# Model settings
OLLAMA_MODEL = "phi3:mini"

OUTPUT_LOG_FILE = "rag_output.txt"
DOC_OUTPUT_FILE = "project_documentation.docx"

SIMILARITY_THRESHOLD = 0.4


# ==============================
# LOAD PROJECT DATASET
# ==============================

def load_projects():

    print("Loading project dataset...", flush=True)

    if not os.path.exists(DATA_PATH):
        raise FileNotFoundError(f"{DATA_PATH} not found")

    with open(DATA_PATH, "r", encoding="utf-8") as f:
        data = json.load(f)

    documents = []

    for item in data:

        if item["documentation"] != "FAILED":

            documents.append(
                Document(
                    page_content=item["documentation"],
                    metadata={"title": item["title"]}
                )
            )

    print("Total valid documents:", len(documents), flush=True)

    return documents


# ==============================
# CREATE CHUNKS
# ==============================

def create_chunks(documents):

    print("Splitting documents...", flush=True)

    splitter = RecursiveCharacterTextSplitter(
        chunk_size=400,
        chunk_overlap=60
    )

    chunks = splitter.split_documents(documents)

    print("Chunks created:", len(chunks), flush=True)

    return chunks


# ==============================
# LOAD EMBEDDINGS
# ==============================

def get_embeddings():

    print("Loading embedding model...", flush=True)

    return HuggingFaceEmbeddings(
        model_name="sentence-transformers/all-MiniLM-L6-v2"
    )


# ==============================
# BUILD VECTOR DATABASE
# ==============================

def build_vector_db(chunks, embeddings):

    print("Creating vector database...", flush=True)

    vector_db = FAISS.from_documents(chunks, embeddings)

    vector_db.save_local(VECTOR_DB_PATH)

    print("Vector database saved!")

    return vector_db


# ==============================
# LOAD VECTOR DATABASE
# ==============================

def load_vector_db(embeddings):

    print("Loading existing vector database...", flush=True)

    return FAISS.load_local(
        VECTOR_DB_PATH,
        embeddings,
        allow_dangerous_deserialization=True
    )


# ==============================
# LOAD LLM
# ==============================

def load_llm():

    print("Loading Ollama model (phi3:mini)...", flush=True)

    return Ollama(
        model=OLLAMA_MODEL,
        base_url="http://localhost:11434",
        num_predict=1200,
        temperature=0
    )


# ==============================
# ASK QUESTION
# ==============================

def ask_question(query, vector_db, llm):

    print("Searching knowledge base...", flush=True)

    results = vector_db.similarity_search_with_score(query, k=3)

    filtered_docs = []

    for doc, score in results:
        if score < SIMILARITY_THRESHOLD:
            filtered_docs.append(doc)

    docs = filtered_docs[:1]

    if len(docs) == 0:
        docs = [doc for doc, score in results[:1]]

    context = docs[0].page_content if docs else ""

    prompt = f"""
You are a professional software project documentation generator.

IMPORTANT RULES:
- The project title MUST be exactly: {query}
- Do NOT change the project name.
- Do NOT introduce other project ideas.
- Do NOT skip any section.
- Implementation steps MUST contain exactly 10 steps.
- Each step must be a SHORT sentence.
- If information is missing, generate reasonable content.

Use the following structure EXACTLY.

Project Title: {query}

Project Overview:
Write 3-4 sentences explaining the project.

Objective:
Write 2-3 sentences describing the goal.

Domain:
Mention the domain clearly.

Software Requirements:
- item
- item
- item

Hardware Requirements:
- item
- item
- item

Workflow:
Explain system workflow briefly.

System Architecture:
Explain architecture briefly.

Input:
Describe system inputs.

Output:
Describe system outputs.

Implementation Steps:
1. Step one.
2. Step two.
3. Step three.
4. Step four.
5. Step five.
6. Step six.
7. Step seven.
8. Step eight.
9. Step nine.
10. Step ten.

Benefits:
- Benefit one
- Benefit two
- Benefit three
- Benefit four
- Benefit five

Future Scope:
- Future improvement one
- Future improvement two
- Future improvement three
- Future improvement four
- Future improvement five

Context information (reference only):
{context}
"""

    return llm.invoke(prompt)


# ==============================
# SAVE TEXT OUTPUT
# ==============================

def save_output(query, answer):

    with open(OUTPUT_LOG_FILE, "w", encoding="utf-8") as f:

        f.write("=" * 60 + "\n")
        f.write(f"Time: {datetime.now()}\n\n")

        f.write("Question:\n")
        f.write(query + "\n\n")

        f.write("Answer:\n")
        f.write(str(answer) + "\n")

        f.write("=" * 60 + "\n")

    print("Output saved ->", OUTPUT_LOG_FILE)


# ==============================
# SAVE WORD DOCUMENT
# ==============================

def save_output_doc(query, answer):

    doc = WordDocument()

    doc.add_heading("Software Project Documentation", level=0)

    doc.add_heading("Project Query", level=1)
    doc.add_paragraph(query)

    doc.add_heading("Generated Documentation", level=1)
    doc.add_paragraph(str(answer))

    doc.save(DOC_OUTPUT_FILE)

    print("Word document created:", DOC_OUTPUT_FILE)


# ==============================
# MAIN
# ==============================

if __name__ == "__main__":

    print("\n===== PROJECT RAG SYSTEM =====\n")

    embeddings = get_embeddings()

    try:

        vector_db = load_vector_db(embeddings)

        print("Using existing vector database")

    except Exception:

        print("Building new vector database...\n")

        documents = load_projects()

        chunks = create_chunks(documents)

        vector_db = build_vector_db(chunks, embeddings)

    llm = load_llm()

    print("RAG system ready\n")

    query = sys.argv[1] if len(sys.argv) > 1 else ""

    if query == "":
        query = input("Enter project topic: ")

    answer = ask_question(query, vector_db, llm)

    print("\nAnswer:\n", answer)

    save_output(query, answer)

    save_output_doc(query, answer)